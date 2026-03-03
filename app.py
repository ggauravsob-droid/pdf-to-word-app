import streamlit as st
import pytesseract
from pdf2image import convert_from_path
from docx import Document
import os
import re
import tempfile
import time  # Naya: Time calculate karne ke liye
# Naya: as_completed ko import kiya taaki live progress dikhe
from concurrent.futures import ThreadPoolExecutor, as_completed

st.set_page_config(page_title="Fast PDF to Word", page_icon="⚡")
st.title("⚡ Fast Scanned PDF Converter")
st.write("Apni PDF upload karein. Yeh app Hindi/English text tezi se extract karega bina extra spaces ke.")

def remove_extra_spaces_and_clean(text):
    cleaned = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', '', text)
    cleaned = re.sub(r'\n\s*\n', '\n\n', cleaned)
    cleaned = re.sub(r' {2,}', ' ', cleaned)
    return cleaned.strip()

def process_single_page(img_path):
    custom_config = r'--oem 3 --psm 6'
    raw_text = pytesseract.image_to_string(img_path, lang='eng+hin', config=custom_config)
    return remove_extra_spaces_and_clean(raw_text)

uploaded_file = st.file_uploader("Apni PDF file yahan drag & drop karein", type=["pdf"])

if uploaded_file is not None:
    st.info("File upload ho gayi hai. Conversion shuru karne ke liye niche button dabayein.")
    
    if st.button("Start Fast Conversion 🚀"):
        with st.spinner("⚡ Initializing... Kripya pratiksha karein"):
            try:
                with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_pdf:
                    temp_pdf.write(uploaded_file.read())
                    pdf_path = temp_pdf.name
                
                word_path = pdf_path.replace('.pdf', '.docx')
                doc = Document()
                
                with tempfile.TemporaryDirectory() as temp_dir:
                    st.write("Images extract ki ja rahi hain... (Isme thoda waqt lag sakta hai)")
                    
                    image_paths = convert_from_path(
                        pdf_path, dpi=150, output_folder=temp_dir, paths_only=True, fmt='jpeg', grayscale=True
                    )
                    
                    total_pages = len(image_paths)
                    st.write(f"Total {total_pages} pages mile. Live OCR shuru ho raha hai...")
                    
                    # LIVE TRACKING UI ELEMENTS
                    progress_bar = st.progress(0)
                    status_text = st.empty()  # Page count dikhane ke liye
                    timer_text = st.empty()   # Time dikhane ke liye
                    
                    start_time = time.time()
                    
                    # Pages ka order sahi rakhne ke liye ek khali list banayi hai
                    extracted_texts = [""] * total_pages 
                    
                    with ThreadPoolExecutor(max_workers=4) as executor:
                        # Har page ko processing ke liye bhej rahe hain
                        future_to_page = {executor.submit(process_single_page, img): i for i, img in enumerate(image_paths)}
                        
                        completed_pages = 0
                        # Jaise-jaise ek-ek page complete hoga, yeh loop chalega
                        for future in as_completed(future_to_page):
                            page_idx = future_to_page[future]
                            try:
                                text = future.result()
                                extracted_texts[page_idx] = text # Sahi sequence mein text save kiya
                            except Exception as e:
                                extracted_texts[page_idx] = f"Error in page {page_idx + 1}"
                            
                            # Update Counters
                            completed_pages += 1
                            elapsed_time = int(time.time() - start_time)
                            
                            # Update Screen (Live UI)
                            progress_bar.progress(completed_pages / total_pages)
                            status_text.success(f"📖 Processed: **{completed_pages} / {total_pages}** pages")
                            timer_text.info(f"⏱️ Time Elapsed: **{elapsed_time} seconds**")
                    
                    # Saara text Word file mein daalna
                    for i, text in enumerate(extracted_texts):
                        if text.strip():
                            doc.add_paragraph(text)
                            doc.add_paragraph("") 
                            
                total_time = int(time.time() - start_time)
                doc.save(word_path)
                st.success(f"✅ Conversion poora ho gaya! Kul samay laga: {total_time} seconds.")
                
                with open(word_path, "rb") as file:
                    st.download_button(
                        label="📥 Download Word File",
                        data=file,
                        file_name=uploaded_file.name.replace('.pdf', '_Fast_Converted.docx'),
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                    
            except Exception as e:
                st.error(f"Ek error aa gaya: {e}")


